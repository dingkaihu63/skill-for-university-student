#!/usr/bin/env python3
"""
思政论文 Word 文档生成器
功能：
  1. 解析用户提供的论文模板/结构要求
  2. 根据模板结构生成规范 Word 文档
  3. 支持自定义封面、标题、正文、参考文献等段落样式
  4. 默认将文件生成至用户桌面
"""

import json
import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

# 修改：无模板时，一律强制宋体（SimSun），小四（12pt）
DEFAULT_TEMPLATE = {
    "page": {
        "top_margin": 2.54,
        "bottom_margin": 2.54,
        "left_margin": 3.17,
        "right_margin": 3.17
    },
    "styles": {
        "cover_title": {"font_name": "SimSun", "font_size": 12, "bold": True, "alignment": "center", "space_after": 24,
                        "color": [0, 0, 0]},
        "cover_info": {"font_name": "SimSun", "font_size": 12, "bold": False, "alignment": "center", "space_after": 12,
                       "color": [0, 0, 0]},
        "heading1": {"font_name": "SimSun", "font_size": 12, "bold": True, "alignment": "left", "space_before": 24,
                     "space_after": 12, "color": [0, 0, 0]},
        "heading2": {"font_name": "SimSun", "font_size": 12, "bold": True, "alignment": "left", "space_before": 18,
                     "space_after": 8, "color": [0, 0, 0]},
        "body": {"font_name": "SimSun", "font_size": 12, "bold": False, "alignment": "justify", "line_spacing": 1.5,
                 "first_line_indent": 2, "space_after": 6, "color": [0, 0, 0]},
        "reference_title": {"font_name": "SimSun", "font_size": 12, "bold": True, "alignment": "left",
                            "space_before": 24, "space_after": 12, "color": [0, 0, 0]},
        "reference_item": {"font_name": "SimSun", "font_size": 12, "bold": False, "alignment": "left",
                           "line_spacing": 1.25, "space_after": 4, "color": [0, 0, 0]}
    },
    "structure": [
        {"type": "cover", "fields": ["title", "author", "school", "date"]},
        {"type": "heading1", "text": "摘要"},
        {"type": "body", "text": "abstract"},
        {"type": "heading1", "text": "关键词"},
        {"type": "body", "text": "keywords"},
        {"type": "heading1", "text": "正文"},
        {"type": "heading1", "text": "参考文献"}
    ]
}


def load_template(template_path=None):
    if template_path and Path(template_path).exists():
        with open(template_path, "r", encoding="utf-8") as f:
            return json.load(f)
    return DEFAULT_TEMPLATE


def set_chinese_font(run, font_name, font_size, bold=False, color=None):
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def create_paragraph(doc, text, style_config, template):
    styles = template.get("styles", {})
    cfg = styles.get(style_config, styles.get("body", {}))
    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    p = doc.add_paragraph()
    p.alignment = alignment_map.get(cfg.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
    pf = p.paragraph_format
    if cfg.get("space_before"):
        pf.space_before = Pt(cfg["space_before"])
    if cfg.get("space_after"):
        pf.space_after = Pt(cfg["space_after"])
    if cfg.get("line_spacing"):
        pf.line_spacing = cfg["line_spacing"]
    if cfg.get("first_line_indent"):
        pf.first_line_indent = Pt(cfg["font_size"] * cfg["first_line_indent"])

    # 按照段落切分，处理换行符
    for line in str(text).split('\n'):
        if line.strip():
            run = p.add_run(line.strip())
            set_chinese_font(run, cfg.get("font_name", "SimSun"), cfg.get("font_size", 12), cfg.get("bold", False),
                             cfg.get("color", [0, 0, 0]))
            run.add_break()
    return p


def generate_docx(content, template, output_path):
    doc = Document()
    page = template.get("page", {})
    for section in doc.sections:
        section.top_margin = Cm(page.get("top_margin", 2.54))
        section.bottom_margin = Cm(page.get("bottom_margin", 2.54))
        section.left_margin = Cm(page.get("left_margin", 3.17))
        section.right_margin = Cm(page.get("right_margin", 3.17))

    structure = template.get("structure", [])
    for section in structure:
        sec_type = section.get("type", "body")
        if sec_type == "cover":
            for field in section.get("fields", []):
                text = content.get(field, f"【{field}】")
                style = "cover_title" if field == "title" else "cover_info"
                create_paragraph(doc, text, style, template)
            doc.add_page_break()
        elif sec_type in ("heading1", "heading2"):
            text = content.get(section.get("text", ""), section.get("text", ""))
            create_paragraph(doc, text, sec_type, template)
        elif sec_type == "body":
            text = content.get(section.get("text", ""), "")
            if text:
                create_paragraph(doc, text, "body", template)
        elif sec_type == "reference_item":
            text = content.get(section.get("text", ""), "")
            if text:
                create_paragraph(doc, text, "reference_item", template)

    # 修改：强制输出至桌面逻辑
    desktop_dir = Path.home() / "Desktop"
    out_file = Path(output_path)
    if not out_file.is_absolute():
        final_path = desktop_dir / out_file.name
    else:
        final_path = out_file

    doc.save(final_path)
    return final_path


def main():
    if len(sys.argv) < 3:
        sys.exit(1)

    content_path = sys.argv[1]
    output_path = sys.argv[2]
    template_path = sys.argv[3] if len(sys.argv) > 3 else None

    with open(content_path, "r", encoding="utf-8") as f:
        content = json.load(f)

    template = load_template(template_path)
    result = generate_docx(content, template, output_path)
    print(f"{result}")


if __name__ == "__main__":
    main()